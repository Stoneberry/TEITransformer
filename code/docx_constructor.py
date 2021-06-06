from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from tei_auxiliary import check_if_rgb
from setup import DOCX_CFG


class DocxConstructor:

    """
    Adapter for working with the python-docx library.
    Contains functions that converts CSS rules to docx functions.
    """

    def __init__(self):
        """
        Initializes the docx object
        """
        self.docx = Document()
        self.paragraph_format = self.docx.styles['Normal'].paragraph_format
        self.paragraph_format.space_after = Pt(5)

    def add_line(self, text):
        """
        Adds new line to the document.
        :param text: str
        :return: python-docx paragraph params (paragraph, run, run.font)
        """
        p = self.docx.add_paragraph()
        run = p.add_run(text)
        font = run.font
        return p, run, font

    @staticmethod
    def add_text_align(p, run, font, value):
        """
        Aligns the text if the specified value is valid.
        :param p: python-docx parameter
        :param run: python-docx parameter
        :param font: python-docx parameter
        :param value: str
        :return: None
        """
        member_index = WD_ALIGN_PARAGRAPH._xml_to_member.get(value)
        if member_index:
            p.alignment = member_index

    @staticmethod
    def add_text_indent(p, run, font, value, measure=""):
        """
        Adds text indent if the measure is valid.
        Only Inches, Pt, Cm measure are allowable.
        :param p: python-docx parameter
        :param run: python-docx parameter
        :param font: python-docx parameter
        :param value: str
        :param measure: str
        :return: None
        """
        mfunc = DOCX_CFG['MEASURES'].get(measure)
        if mfunc:
            p.paragraph_format.first_line_indent = globals()[mfunc](value)

    @staticmethod
    def add_margin_left(p, run, font, value, measure=""):
        """
        Adds left margin if the measure is valid.
        Only Inches, Pt, Cm measure are allowable.
        :param p: python-docx parameter
        :param run: python-docx parameter
        :param font: python-docx parameter
        :param value: str
        :param measure: str
        :return: None
        """
        mfunc = DOCX_CFG['MEASURES'].get(measure)
        if mfunc:
            p.paragraph_format.left_indent = globals()[mfunc](value)

    @staticmethod
    def add_margin_right(p, run, font, value, measure=""):
        """
        Adds right margin if the measure is valid.
        Only Inches, Pt, Cm measure are allowable.
        :param p: python-docx parameter
        :param run: python-docx parameter
        :param font: python-docx parameter
        :param value: str
        :param measure: str
        :return: None
        """
        mfunc = DOCX_CFG['MEASURES'].get(measure)
        if mfunc:
            p.paragraph_format.right_indent = globals()[mfunc](value)

    @staticmethod
    def add_margin_bottom(p, run, font, value, measure=""):
        """
        Adds bottom margin if the measure is valid.
        Only Inches, Pt, Cm measure are allowable.
        :param p: python-docx parameter
        :param run: python-docx parameter
        :param font: python-docx parameter
        :param value: str
        :param measure: str
        :return: None
        """
        mfunc = DOCX_CFG['MEASURES'].get(measure)
        if mfunc:
            p.paragraph_format.space_after = globals()[mfunc](value)

    @staticmethod
    def add_margin_top(p, run, font, value, measure=""):
        """
        Adds top margin if the measure is valid.
        Only Inches, Pt, Cm measure are allowable.
        :param p: python-docx parameter
        :param run: python-docx parameter
        :param font: python-docx parameter
        :param value: str
        :param measure: str
        :return: None
        """
        mfunc = DOCX_CFG['MEASURES'].get(measure)
        if mfunc:
            p.paragraph_format.space_before = globals()[mfunc](value)

    @staticmethod
    def add_line_height(p, run, font, value, measure=""):
        """
        Adds line height if the measure is valid.
        Only Inches, Pt, Cm measure are allowable.
        :param p: python-docx parameter
        :param run: python-docx parameter
        :param font: python-docx parameter
        :param value: str
        :param measure: str
        :return: None
        """
        mfunc = DOCX_CFG['MEASURES'].get(measure)
        if mfunc:
            p.paragraph_format.line_spacing = globals()[mfunc](value)

    @staticmethod
    def add_font_weight(p, run, font, value):
        """
        Changes font weight if the specified value is valid.
        Only bold, normal are allowable.
        :param p: python-docx parameter
        :param run: python-docx parameter
        :param font: python-docx parameter
        :param value: str
        :return: None
        """
        params = DOCX_CFG['FONT_VALUES']['weight'].get(value)
        if params:
            font.__setattr__(*params)

    @staticmethod
    def add_font_style(p, run, font, value):
        """
        Changes font style if the specified value is valid.
        Only italic, normal are allowable.
        :param p: python-docx parameter
        :param run: python-docx parameter
        :param font: python-docx parameter
        :param value: str
        :return: None
        """
        params = DOCX_CFG['FONT_VALUES']['style'].get(value)
        if params:
            font.__setattr__(*params)

    @staticmethod
    def add_color(p, run, font, value):
        """
        Changes font color if the specified value is valid.
        Only RGB notation is valid.
        :param p: python-docx parameter
        :param run: python-docx parameter
        :param font: python-docx parameter
        :param value: str
        :return: None
        """
        if check_if_rgb(value):
            font.color.rgb = RGBColor(*value)

    @staticmethod
    def add_font_size(p, run, font, value, measure=""):
        """
        Changes font size if the measure is valid.
        Only Inches, Pt, Cm measure are allowable.
        :param p: python-docx parameter
        :param run: python-docx parameter
        :param font: python-docx parameter
        :param value: str
        :param measure: str
        :return: None
        """
        mfunc = DOCX_CFG['MEASURES'].get(measure)
        if mfunc:
            font.size = globals()[mfunc](value)

    @staticmethod
    def add_text_transform(p, run, font, value):
        """
        Changes text capitalization if the specified value is valid.
        Only uppercase and lowercase are valid.
        :param p: python-docx parameter
        :param run: python-docx parameter
        :param font: python-docx parameter
        :param value: str
        :return: None
        """
        params = DOCX_CFG['FONT_VALUES']['transform'].get(value)
        if params:
            font.__setattr__(*params)

    @staticmethod
    def add_text_decoration(p, run, font, value):
        """
        Changes text decoration if the specified value is valid.
        Only underline and normal are valid.
        :param p: python-docx parameter
        :param run: python-docx parameter
        :param font: python-docx parameter
        :param value: str
        :return: None
        """
        params = DOCX_CFG['FONT_VALUES']['decoration'].get(value)
        if params:
            font.__setattr__(*params)

    @staticmethod
    def add_pb(p, run, font):
        """
        Adds page break.
        :param p: python-docx parameter
        :param run: python-docx parameter
        :param font: python-docx parameter
        :return: None
        """
        run.add_break(WD_BREAK.PAGE)
        # p.paragraph_format.page_break_before = True

    def save(self, output_filename):
        """
        Saves the file
        :param output_filename: str
        :return: None
        """
        self.docx.save(output_filename)
